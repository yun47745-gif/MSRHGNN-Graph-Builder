from docx import Document
from docx.shared import Pt

def create_paper_doc():
    doc = Document()
    
    # 标题
    title = doc.add_heading('2. Dataset (数据集与网络构建)', level=1)
    
    # 2.1
    doc.add_heading('2.1 数据来源', level=2)
    doc.add_paragraph('为了全面揭示神经精神类疾病的潜在病理机制与药物作用靶点，本研究整合了多个具有高置信度的权威生物医学数据库。首先，从比较毒理基因组学数据库 (CTD) 中，通过检索 Alzheimer、Schizophrenia、Depression 等神经精神类疾病关键词，筛选出 281 种核心疾病实体。基于 CTD 的人工审阅（Curated）数据，我们提取了具备直接实验证据（Marker 或 Mechanism）的疾病-基因强关联数据，以及药物-疾病治疗关联数据。')
    doc.add_paragraph('在药物与分子层面，本研究引入了 DrugBank 数据库。我们解析了其官方 XML 数据，提取了近两万种药物的化学结构（SMILES 序列）以及对应的靶点蛋白（Targets）。为了消除异构数据库间的标识符壁垒，我们实施了严格的多级 ID 对齐策略：通过 MeSH ID 与药物名称的双重校验对齐了 CTD 与 DrugBank 的药物实体；同时利用 UniProtKB 提供的官方映射表，将靶点蛋白的 UniProt Accession 统一转换为 Entrez Gene ID，从而实现了与 CTD 基因实体的精准锚定。')
    doc.add_paragraph('此外，为了构建基因互作网络与疾病临床表型网络，我们分别引入了 STRING (v12.0) 数据库中仅限人类（Taxonomy ID: 9606）的蛋白质-蛋白质互作（PPI）数据，以及人类表型本体数据库 (HPO) 的表型注释（phenotype.hpoa）与层级拓扑（hp.obo）数据。')

    # 2.2
    doc.add_heading('2.2 多源相似性网络构建', level=2)
    doc.add_paragraph('传统的单一相似性网络往往存在严重的“信息孤岛”与数据稀疏问题。为此，本研究从“宏观临床表型”与“微观基因机制”两个维度，构建了双重视角的疾病高阶相似性网络：')
    
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('临床表型相似性网络 (H视图, Sim_H)：').bold = True
    p1.add_run('基于 HPO 数据库，我们首先统计了各个表型术语的注释频率，并计算其信息量（Information Content, IC）。对于任意两种疾病，通过提取它们在 HPO 有向无环图（DAG）中的共同祖先节点，并寻找这些公共节点中的最大 IC 值，计算出两种疾病间的表型相似度。该视图能够有效量化疾病在宏观临床症状上的关联程度。')
    
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('基因机制相似性网络 (G视图, Sim_G)：').bold = True
    p2.add_run('疾病的发生往往由一组相互作用的基因模块驱动。我们过滤了 STRING 网络中置信度得分低于 700 的噪声边，保留了高质量的 PPI 拓扑结构。对于任意两种疾病，计算其关联基因集在 PPI 网络中的相互作用得分累加期望。特别地，当不同疾病共享相同的致病基因时，赋予最高互作权重。该视图从底层分子互作层面揭示了疾病间的潜在机制连通性。')
    
    doc.add_paragraph('为避免密集网络引入过多的拓扑噪声并降低计算复杂度，本研究对上述两个相似度矩阵进行了 k-近邻 (k-NN) 稀疏化处理（本实验设置 k=5），仅保留与中心节点最具生物学相关性的顶级关联边。')

    # 2.3
    doc.add_heading('2.3 异构生物网络构建', level=2)
    doc.add_paragraph('基于上述提取的实体特征与关系拓扑，本研究构建了一个统一的异构图属性网络（Heterogeneous Graph）。该网络被封装为 PyTorch Geometric 的数据结构，具体包含以下组件：')
    
    doc.add_paragraph('1. 节点与特征 (Nodes & Features)：包含药物 (Drug)、疾病 (Disease) 与基因 (Gene) 三类节点。为了保留化学语义信息，我们利用 RDKit 解析药物的 SMILES 序列，生成了 1024 维的 Morgan 分子指纹（Morgan Fingerprints）作为药物节点的初始特征向量；疾病与基因节点则预留了对应维度的语义与表征向量空间。', style='List Number')
    doc.add_paragraph('2. 异构关系边 (Heterogeneous Edges)：网络中构建了跨越不同实体类型的物理交互边，包括药物治疗疾病边 (drug, treats, disease) 以及疾病关联基因边 (disease, associated_with, gene)。', style='List Number')
    doc.add_paragraph('3. 同构拓扑边 (Homogeneous Edges)：将 2.2 节中构建的两种相似度网络转化为疾病节点间的拓扑连接，即临床相似度边 (disease, sim_h, disease) 与机制相似度边 (disease, sim_g, disease)，并将其相似度得分作为边权重注入模型。', style='List Number')
    
    doc.add_paragraph('最终构建的异构图不仅建立了显式的“药物-疾病-基因”闭环链路，还通过双重相似度视图打破了疾病间的孤立状态，为后续图神经网络的跨模态消息传递提供了丰富的结构化先验知识。')

    # References
    doc.add_heading('References (参考文献)', level=1)
    refs = [
        "[1] Davis, A. P., et al. (2021). The Comparative Toxicogenomics Database: update 2021. Nucleic Acids Research, 49(D1), D1138-D1143.",
        "[2] Wishart, D. S., et al. (2018). DrugBank 5.0: a major update to the DrugBank database for 2018. Nucleic Acids Research, 46(D1), D1074-D1082.",
        "[3] Szklarczyk, D., et al. (2023). The STRING database in 2023: protein–protein association networks and functional enrichment analyses for any sequenced genome of interest. Nucleic Acids Research, 51(D1), D638-D646.",
        "[4] Köhler, S., et al. (2021). The Human Phenotype Ontology in 2021. Nucleic Acids Research, 49(D1), D1207-D1217.",
        "[5] Rogers, D., & Hahn, M. (2010). Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 50(5), 742-754.",
        "[6] Fey, M., & Lenssen, J. E. (2019). Fast graph representation learning with PyTorch Geometric. arXiv preprint arXiv:1903.02428.",
        "[7] Wang, X., et al. (2019). Heterogeneous graph attention network. In The World Wide Web Conference (pp. 2022-2032)."
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # 保存文件
    doc.save('Dataset_and_Network_Construction.docx')
    print("✅ 文档已成功生成：Dataset_and_Network_Construction.docx")

if __name__ == '__main__':
    create_paper_doc()